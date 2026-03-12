from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "input_templates" / "resume_input_template_word.docx"


def main() -> None:
    doc = Document()
    doc.add_heading("中文简历填写模板", level=0)
    doc.add_paragraph("姓名：XXX")
    doc.add_paragraph("求职意向：XXX")
    doc.add_paragraph("基本信息：手机号 | 邮箱 | GitHub | 城市 | 其他信息")

    sections = [
        (
            "教育经历",
            [
                "1. 学校名称 | 专业名称 学位名称 | 20XX.XX -- 20XX.XX",
                "- GPA：X.X/X.X（排名X/X）。",
                "- 奖项、科研、论文、专利、软著等信息。",
                "2. 学校名称 | 专业名称 学位名称 | 20XX.XX -- 20XX.XX",
                "- GPA：X.X/X.X（排名X/X）。",
                "- 主修方向、奖项或实践信息。",
            ],
        ),
        (
            "工作经历",
            [
                "1. 公司名称 | 岗位名称 | 20XX.XX -- 20XX.XX",
                "- title_link: https://example.com/experience-a",
                "- 负责的业务、模块或产品方向。",
                "- 你的核心动作：调研、需求分析、方案设计、协作推进、数据分析等。",
                "- 量化结果：效率、转化、准确率、成本、满意度等。",
                "  - 补充说明：用于放更细一层的动作或结果。",
                "    - 更细一层：用于特别关键的补充信息。",
                "2. 公司名称 | 岗位名称 | 20XX.XX -- 20XX.XX",
                "- 负责的业务、模块或产品方向。",
                "- 你的核心动作。",
                "- 量化结果。",
            ],
        ),
        (
            "项目经历",
            [
                "1. 项目名称 | 角色 | 20XX.XX -- 20XX.XX",
                "- title_link: https://github.com/yourname/project-a",
                "- 问题背景：这个项目要解决什么问题。",
                "- 方案设计：你如何定义产品或技术方案。",
                "- 落地过程：你做了哪些关键动作。",
                "  - 子项说明：如果需要展开关键模块，可继续缩进。",
                "- 结果验证：指标、反馈、奖项或上线结果。",
                "2. 项目名称 | 角色 | 20XX.XX -- 20XX.XX",
                "- 问题背景。",
                "- 方案设计。",
                "- 落地过程。",
                "- 结果验证。",
            ],
        ),
        (
            "校园经历",
            [
                "1. 组织名称 | 职务 | 20XX.XX -- 20XX.XX",
                "- 负责事项。",
                "- 协作或组织动作。",
                "- 结果或影响。",
                "2. 组织名称 | 职务 | 20XX.XX -- 20XX.XX",
                "- 负责事项。",
                "- 协作或组织动作。",
                "- 结果或影响。",
            ],
        ),
        (
            "技能",
            [
                "产品能力：XXX",
                "AI 应用：XXX",
                "工具与表达：XXX",
            ],
        ),
    ]

    for title, paragraphs in sections:
        doc.add_paragraph(f"【{title}】")
        for text in paragraphs:
            doc.add_paragraph(text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
